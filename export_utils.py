from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, PageBreak, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, mm
from reportlab.lib.pagesizes import A4 # Using A4 as per metric preference
from reportlab.platypus import HRFlowable, KeepTogether
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.lib.colors import black

from models import CVData, PersonalInformation, Education, Experience, Skills # Import CVData and its components

# Helper function for drawing a line
def draw_line(canvas, doc):
    canvas.line(doc.leftMargin, doc.y, doc.width + doc.leftMargin, doc.y)

def export_to_pdf(cv_data: CVData, filename: str) -> BytesIO:
    buffer = BytesIO()
    
    # Use A4 and define margins in mm
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=20*mm, leftMargin=20*mm,
                            topMargin=20*mm, bottomMargin=20*mm)
    
    styles = getSampleStyleSheet()
    
    # Custom styles based on a professional template (using metric units for sizes)
    styles.add(ParagraphStyle(name='NameStyle',
                             fontName='Helvetica-Bold',
                             fontSize=28,
                             leading=32,
                             alignment=TA_CENTER,
                             spaceAfter=2*mm))

    styles.add(ParagraphStyle(name='ContactInfoStyle',
                             fontName='Helvetica',
                             fontSize=10,
                             leading=12,
                             alignment=TA_CENTER,
                             spaceAfter=8*mm, # Increased space for separation
                             textColor=black))

    styles.add(ParagraphStyle(name='SectionTitleStyle',
                             fontName='Helvetica-Bold',
                             fontSize=16,
                             leading=18,
                             alignment=TA_LEFT,
                             spaceBefore=10*mm, # Space before section title
                             spaceAfter=4*mm,
                             textColor=black))
    
    styles.add(ParagraphStyle(name='SubHeadingStyleBold',
                             fontName='Helvetica-Bold',
                             fontSize=12,
                             leading=14,
                             alignment=TA_LEFT,
                             spaceAfter=1*mm,
                             textColor=black))
    
    styles.add(ParagraphStyle(name='SubHeadingStyleNormal',
                             fontName='Helvetica',
                             fontSize=12,
                             leading=14,
                             alignment=TA_LEFT,
                             spaceAfter=1*mm,
                             textColor=black))

    styles.add(ParagraphStyle(name='DateLocationStyle',
                             fontName='Helvetica',
                             fontSize=10,
                             leading=12,
                             alignment=TA_LEFT, # Changed to left for now, will adjust based on template
                             spaceAfter=4*mm,
                             textColor=black))

    styles.add(ParagraphStyle(name='NormalTextStyle',
                             fontName='Helvetica',
                             fontSize=10,
                             leading=12,
                             alignment=TA_LEFT,
                             spaceAfter=4*mm,
                             textColor=black))
    
    styles.add(ParagraphStyle(name='BulletPointStyle',
                             fontName='Helvetica',
                             fontSize=10,
                             leading=12,
                             alignment=TA_LEFT,
                             leftIndent=5*mm,
                             bulletIndent=2*mm,
                             bulletText='•', # Unicode for a standard bullet point
                             spaceBefore=0.5*mm,
                             spaceAfter=0.5*mm,
                             textColor=black))

    story = []

    # --- Personal Information ---
    if cv_data.personal_info.name:
        story.append(Paragraph(cv_data.personal_info.name, styles['NameStyle']))
        
        contact_details = []
        if cv_data.personal_info.email:
            contact_details.append(cv_data.personal_info.email)
        if cv_data.personal_info.phone:
            contact_details.append(cv_data.personal_info.phone)
        if cv_data.personal_info.linkedin:
            contact_details.append(cv_data.personal_info.linkedin) # Just the URL, not "LinkedIn: "
        if cv_data.personal_info.github:
            contact_details.append(cv_data.personal_info.github) # Just the URL, not "GitHub: "
        
        if contact_details:
            story.append(Paragraph(" | ".join(contact_details), styles['ContactInfoStyle']))
        
    # --- Summary ---
    if cv_data.personal_info.summary:
        story.append(Paragraph("SUMMARY", styles['SectionTitleStyle']))
        story.append(HRFlowable(width="100%", thickness=1, color=black)) # Horizontal line separator
        story.append(Spacer(1, 2*mm))
        story.append(Paragraph(cv_data.personal_info.summary, styles['NormalTextStyle']))
        story.append(Spacer(1, 4*mm)) # Space after summary

    # --- Education ---
    if cv_data.education:
        story.append(Paragraph("EDUCATION", styles['SectionTitleStyle']))
        story.append(HRFlowable(width="100%", thickness=1, color=black)) # Horizontal line separator
        story.append(Spacer(1, 2*mm))
        for edu in cv_data.education:
            edu_details = []
            edu_details.append(Paragraph(f"<b>{edu.degree}</b> in {edu.major}", styles['SubHeadingStyleBold']))
            edu_details.append(Paragraph(f"{edu.institution}, {edu.location}", styles['SubHeadingStyleNormal']))
            
            date_gpa_info = []
            if edu.start_date and edu.end_date:
                date_gpa_info.append(f"{edu.start_date} - {edu.end_date}")
            if edu.gpa:
                date_gpa_info.append(f"GPA: {edu.gpa}")
            
            if date_gpa_info:
                edu_details.append(Paragraph(" | ".join(date_gpa_info), styles['DateLocationStyle']))
            
            story.append(KeepTogether(edu_details)) # Keep education details together
            story.append(Spacer(1, 4*mm)) # Space after each education entry

    # --- Experience ---
    if cv_data.experience:
        story.append(Paragraph("EXPERIENCE", styles['SectionTitleStyle']))
        story.append(HRFlowable(width="100%", thickness=1, color=black)) # Horizontal line separator
        story.append(Spacer(1, 2*mm))
        for exp in cv_data.experience:
            exp_details = []
            exp_details.append(Paragraph(f"<b>{exp.title}</b> at {exp.company}", styles['SubHeadingStyleBold']))
            exp_details.append(Paragraph(f"{exp.location}", styles['SubHeadingStyleNormal']))
            
            if exp.start_date and exp.end_date:
                exp_details.append(Paragraph(f"{exp.start_date} - {exp.end_date}", styles['DateLocationStyle']))
            
            story.append(KeepTogether(exp_details)) # Keep experience header together

            if exp.description:
                # Using ListFlowable for proper bullet indentation
                bullet_list_items = [ListItem(Paragraph(line.strip(), styles['BulletPointStyle'])) for line in exp.description.split('\n') if line.strip()]
                if bullet_list_items:
                    story.append(ListFlowable(bullet_list_items,
                                              bulletType='bullet',
                                              start='bullet',
                                              indent=5*mm, # Indent the whole list
                                              bulletAnchor='start',
                                              spaceBefore=1*mm,
                                              spaceAfter=2*mm))
            story.append(Spacer(1, 4*mm)) # Space after each experience entry

    # --- Skills ---
    if cv_data.skills.technical or cv_data.skills.soft or cv_data.skills.languages:
        story.append(Paragraph("SKILLS", styles['SectionTitleStyle']))
        story.append(HRFlowable(width="100%", thickness=1, color=black)) # Horizontal line separator
        story.append(Spacer(1, 2*mm))
        
        # Using a list for skills to handle multiple types
        skill_items = []
        if cv_data.skills.technical:
            skill_items.append(Paragraph(f"<b>Technical Skills:</b> {', '.join(cv_data.skills.technical)}", styles['NormalTextStyle']))
        if cv_data.skills.soft:
            skill_items.append(Paragraph(f"<b>Soft Skills:</b> {', '.join(cv_data.skills.soft)}", styles['NormalTextStyle']))
        if cv_data.skills.languages:
            skill_items.append(Paragraph(f"<b>Languages:</b> {', '.join(cv_data.skills.languages)}", styles['NormalTextStyle']))
        
        for item in skill_items:
            story.append(item)
        story.append(Spacer(1, 4*mm))


    doc.build(story)
    buffer.seek(0)
    return buffer

def export_to_docx(cv_data: CVData, filename: str) -> BytesIO:
    document = Document()
    
    # Set up basic styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)

    # Personal Information
    if cv_data.personal_info.name:
        heading = document.add_heading(level=1)
        runner = heading.add_run(cv_data.personal_info.name)
        runner.font.size = Pt(24)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact_details = []
        if cv_data.personal_info.email:
            contact_details.append(cv_data.personal_info.email)
        if cv_data.personal_info.phone:
            contact_details.append(cv_data.personal_info.phone)
        if cv_data.personal_info.linkedin:
            contact_details.append(f"LinkedIn: {cv_data.personal_info.linkedin}")
        if cv_data.personal_info.github:
            contact_details.append(f"GitHub: {cv_data.personal_info.github}")
        
        if contact_details:
            contact_para = document.add_paragraph(" | ".join(contact_details))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph() # Add some space


    if cv_data.personal_info.summary:
        document.add_heading("Summary", level=2)
        document.add_paragraph(cv_data.personal_info.summary)
        document.add_paragraph() # Add space

    # Education
    if cv_data.education:
        document.add_heading("Education", level=2)
        for edu in cv_data.education:
            document.add_paragraph(f"<b>{edu.degree}</b> in {edu.major}", style='Intense Quote').runs[0].bold = True # Using Intense Quote for distinct style
            document.add_paragraph(f"{edu.institution}, {edu.location}")
            document.add_paragraph(f"{edu.start_date} - {edu.end_date}")
            if edu.gpa:
                document.add_paragraph(f"GPA: {edu.gpa}")
            document.add_paragraph() # Add space

    # Experience
    if cv_data.experience:
        document.add_heading("Experience", level=2)
        for exp in cv_data.experience:
            document.add_paragraph(f"<b>{exp.title}</b> at {exp.company}, {exp.location}", style='Intense Quote').runs[0].bold = True
            document.add_paragraph(f"{exp.start_date} - {exp.end_date}")
            if exp.description:
                for line in exp.description.split('\n'):
                    if line.strip():
                        document.add_paragraph(line.strip(), style='List Bullet')
            document.add_paragraph() # Add space

    # Skills
    if cv_data.skills.technical or cv_data.skills.soft or cv_data.skills.languages:
        document.add_heading("Skills", level=2)
        if cv_data.skills.technical:
            tech_para = document.add_paragraph()
            tech_para.add_run("Technical Skills: ").bold = True
            tech_para.add_run(', '.join(cv_data.skills.technical))
        if cv_data.skills.soft:
            soft_para = document.add_paragraph()
            soft_para.add_run("Soft Skills: ").bold = True
            soft_para.add_run(', '.join(cv_data.skills.soft))
        if cv_data.skills.languages:
            lang_para = document.add_paragraph()
            lang_para.add_run("Languages: ").bold = True
            lang_para.add_run(', '.join(cv_data.skills.languages))
        document.add_paragraph() # Add space

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer